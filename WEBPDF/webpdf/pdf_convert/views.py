from rest_framework import generics, status
from rest_framework.response import Response
from database.models import Data
from database.serializers import DataSerializer
import ocrmypdf
import os
from PyPDF2 import PdfReader
from django.http import JsonResponse
import requests
from urllib.parse import quote
from docx import Document
from rest_framework.decorators import api_view, permission_classes
from user.views import IsStaff

# Upload PDF View
"""class UploadPDFView(generics.CreateAPIView):
    queryset = PDFFile.objects.all()
    serializer_class = PDFFileSerializer
"""
        
# Convert PDF View (update logic)
class ConvertPDFView(generics.UpdateAPIView):
    queryset = Data.objects.filter(active = True, convert=False, type = 'application/pdf')
    serializer_class = DataSerializer
    #trich xuat van ban
    def extract_text_from_pdf(self, pdf_path):
            """Hàm trích xuất văn bản từ file PDF"""
            try:
                reader = PdfReader(pdf_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
            except Exception as e:
                raise ValueError(f"Lỗi khi trích xuất văn bản: {str(e)}")
            
    def update(self, request, *args, **kwargs):
        pdf_file = self.get_object()
        input_path = pdf_file.original_file.path

        output_filename = f'converted_{pdf_file.name}'
        if not output_filename.endswith('.pdf'):
            output_filename += '.pdf'

        output_path = os.path.join('media/pdfconvert/', output_filename)

        try:
            # Kiem tra va tao thu muc neu chua ton tai
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Run OCR de chuyen doi file
            try:
                ocrmypdf.ocr(input_path, output_path, force_ocr=True, lang="vie")
            except ocrmypdf.exceptions.PriorOcrFoundError:
                return Response({"error": "PDF da co du lieu OCR."}, status=status.HTTP_400_BAD_REQUEST)
            except ocrmypdf.exceptions.MissingDependencyError:
                return Response({"error": "Thieu phu thuoc can thiet cho OCRmyPDF."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            except ocrmypdf.exceptions.InputFileError:
                return Response({"error": "Dinh dang file dau vao khong hop le hoac khong duoc ho tro."}, status=status.HTTP_400_BAD_REQUEST)
            except ocrmypdf.exceptions.OutputFileAccessError:
                return Response({"error": "Khong the ghi file dau ra. Kiem tra quyen truy cap file."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            except ocrmypdf.exceptions.SubprocessOutputError as e:
                return Response({"error": f"Loi xay ra trong qua trinh OCR: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            except Exception as e:
                return Response({"error": f"Loi khong mong muon xay ra trong qua trinh OCR: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Cap nhat file da duoc convert
            try:
                # Trích xuất văn bản từ file đã OCR
                extracted_text = self.extract_text_from_pdf(output_path)

                encoded_file_name = quote(output_filename)
                download_url = f"http://127.0.0.1:8000/database/download_converted/{encoded_file_name}"
                # luu database
                pdf_file.download_converted_file = download_url
                pdf_file.converted_file.name = f'pdfconvert/{output_filename}'
                pdf_file.text_content = extracted_text  # Lưu nội dung vào TextField
                pdf_file.convert = True
                pdf_file.save()
            except Exception as e:
                return Response({"error": f"Loi cap nhat co so du lieu: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            return Response({"message": "Chuyen doi file thanh cong"}, status=status.HTTP_200_OK)

        except FileNotFoundError:
            return Response({"error": "Khong tim thay file dau vao."}, status=status.HTTP_404_NOT_FOUND)
        except PermissionError:
            return Response({"error": "Khong du quyen truy cap file."}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except OSError as e:
            return Response({"error": f"Loi he thong: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            return Response({"error": f"Loi khong mong muon xay ra: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ConvertDocxView(generics.UpdateAPIView):  
    queryset = Data.objects.filter(active=True, convert=False, type__icontains='docx')
    serializer_class = DataSerializer
    def extract_text_from_docx(self, file):
        # Hàm xử lý file DOCX và trả về nội dung dạng chuỗi
        doc = Document(file)
        extracted_text = ""
        # Đọc từng đoạn văn bản
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Bỏ qua các đoạn rỗng
                extracted_text += paragraph.text.strip() + "\n"

        # Đọc các bảng trong file
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                extracted_text += " | ".join(row_data) + "\n"
        return extracted_text
    
    def update(self, request, *args, **kwargs):
        docx_file = self.get_object()
        input_path = docx_file.original_file.path
        extracted_text = self.extract_text_from_docx(input_path)
        docx_file.text_content = extracted_text  # Lưu nội dung vào TextField
        docx_file.convert = True
        docx_file.save()
        return Response({"message": "Luu text file docx thanh cong"}, status=status.HTTP_200_OK)

class ConvertOctetView(generics.UpdateAPIView):
    queryset = Data.objects.filter(active = True, convert=False, type = 'application/octet-stream')
    serializer_class = DataSerializer
    #trich xuat van ban
    def extract_text_from_pdf(self, pdf_path):
            """Hàm trích xuất văn bản từ file PDF"""
            try:
                reader = PdfReader(pdf_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
            except Exception as e:
                raise ValueError(f"Lỗi khi trích xuất văn bản: {str(e)}")
            
    def update(self, request, *args, **kwargs):
        pdf_file = self.get_object()
        input_path = pdf_file.original_file.path
        extracted_text = self.extract_text_from_pdf(input_path)
        pdf_file.text_content = extracted_text  # Lưu nội dung vào TextField
        pdf_file.convert = True
        pdf_file.save()
        return Response({"message": "Luu text file docx thanh cong"}, status=status.HTTP_200_OK)

@api_view(['GET'])
@permission_classes([IsStaff])  # Áp dụng phân quyền
def convert_files(request):
    def process_conversion(data_url, convert_url_template):
        try:
            # Gửi yêu cầu GET để lấy danh sách dữ liệu
            response_get = requests.get(data_url)

            if response_get.status_code == 200:
                data = response_get.json()  # Chuyển đổi dữ liệu JSON thành danh sách Python

                # Duyệt qua từng item và lấy id
                for item in data:
                    id_value = item['id']

                    # URL API thứ hai để thực hiện PUT
                    convert_url = convert_url_template.format(id_value=id_value)

                    # Gửi yêu cầu PUT tới API thứ hai
                    response_put = requests.put(convert_url)

                    # Kiểm tra kết quả của yêu cầu PUT
                    if response_put.status_code == 200:
                        print(f"Convert Data có ID {id_value} thành công")
                    else:
                        print(f"Lỗi khi gửi PUT cho ID {id_value}: {response_put.status_code}")

                return {"status": "success", "message": f"Đã convert xong dữ liệu từ {data_url}"}
            else:
                return {"status": "error", "message": f"Failed to fetch data: {response_get.status_code}"}

        except requests.exceptions.RequestException as e:
            return {"status": "error", "message": str(e)}

    # Xử lý chuỗi logic cho PDF
    pdf_result = process_conversion(
        data_url="http://127.0.0.1:8000/database/data/pdfnotconvert/",
        convert_url_template="http://127.0.0.1:8000/pdf_convert/convertpdf/{id_value}/"
    )
    if pdf_result["status"] == "error":
        return JsonResponse(pdf_result)

    # Xử lý chuỗi logic cho DOCX
    docx_result = process_conversion(
        data_url="http://127.0.0.1:8000/database/data/docxnotconvert/",
        convert_url_template="http://127.0.0.1:8000/pdf_convert/convertdocx/{id_value}/"
    )
    if docx_result["status"] == "error":
        return JsonResponse(docx_result)

    # Xử lý chuỗi logic cho octet-stream
    octet_result = process_conversion(
        data_url="http://127.0.0.1:8000/database/data/octetnotconvert/",
        convert_url_template="http://127.0.0.1:8000/pdf_convert/convertoctet/{id_value}/"
    )
    if octet_result["status"] == "error":
        return JsonResponse(octet_result)

    return JsonResponse({"status": "success", "message": "Tất cả dữ liệu PDF, DOCX và octet-stream đã được convert"})

