from rest_framework import viewsets
from .models import File
from .serializers import FileSerializer
from django.views.decorators.csrf import csrf_exempt
from rest_framework.response import Response
from rest_framework import status
from rest_framework.decorators import api_view
import os
import ocrmypdf
from urllib.parse import quote
from django.conf import settings
import mimetypes
from django.http import HttpResponse, Http404
from rest_framework.decorators import action
from docx import Document
import pytesseract
from pdf2image import convert_from_path

class FileViewSet(viewsets.ModelViewSet):
    queryset = File.objects.all()
    serializer_class = FileSerializer
    @action(detail=False, methods=['delete'], url_path='delete-all')
    def delete_all(self, request):
        # Xóa tất cả các bản ghi trong model File
        File.objects.all().delete()
        return Response({"message": "All files have been deleted."}, status=status.HTTP_204_NO_CONTENT)
    
def download_pdf_file(request, filename):
    # Đường dẫn tới thư mục chứa file
    file_path = os.path.join(settings.MEDIA_ROOT, 'nguoidung/pdf', filename)
    if os.path.exists(file_path):
        mime_type, _ = mimetypes.guess_type(file_path)  # Xác định MIME type
        with open(file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type=mime_type or 'application/octet-stream')
            response['Content-Disposition'] = f'inline; filename="{filename}"'  # Thiết lập inline
            return response
    else:
        raise Http404("File not found.")
    
def download_docx_file(request, filename):
    # Đường dẫn tới thư mục chứa file
    file_path = os.path.join(settings.MEDIA_ROOT, 'nguoidung/doc', filename)
    if os.path.exists(file_path):
        mime_type, _ = mimetypes.guess_type(file_path)  # Xác định MIME type
        with open(file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type=mime_type)
            response['Content-Disposition'] = f'inline; filename="{filename}"'  # Thiết lập inline
            return response
    else:
        raise Http404("File not found.")

@api_view(['POST'])
def upload_and_get_pdf(request):
    if 'file' in request.FILES:
        # Nhận file từ yêu cầu
        uploaded_file = request.FILES['file']
        
        # Tạo một đối tượng File và gắn original_file với file đã upload
        new_file = File(original_file=uploaded_file)
        
        # Lưu đối tượng File vào cơ sở dữ liệu
        new_file.save()
        #convert pdf -> pdf
        input_path = new_file.original_file.path
        output_filename =new_file.name
        output_path = os.path.join('media/nguoidung/pdf/', output_filename)
        
        # Kiem tra va tao thu muc neu chua ton tai
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        ocrmypdf.ocr(input_path, output_path, force_ocr=True, lang="vie")

        encoded_file_name = quote(output_filename)
        download_url = f"http://127.0.0.1:8000/nguoi_dung/download_pdf_converted/{encoded_file_name}"
        new_file.converted_pdf.name = f'nguoidung/pdf/{output_filename}'
        new_file.preview_file = download_url
        new_file.save()
        # Lấy dữ liệu JSON của File vừa tạo bằng serializer
        serializer = FileSerializer(new_file)
        
        # Trả về JSON của File cùng với HTTP status code 201 (Created)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    # Nếu file không được gửi trong yêu cầu
    return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['POST'])
def upload_and_get_docx(request):
    if 'file' in request.FILES:
        # Nhận file từ yêu cầu
        uploaded_file = request.FILES['file']
        
        # Tạo một đối tượng File và gắn original_file với file đã upload
        new_file = File(original_file=uploaded_file)
        
        # Lưu đối tượng File vào cơ sở dữ liệu
        new_file.save()
        
        # Convert PDF → DOCX
        input_path = new_file.original_file.path
        output_filename = os.path.splitext(new_file.name)[0] + '.docx'
        output_path = os.path.join('media/nguoidung/doc/', output_filename)

        # Kiểm tra và tạo thư mục nếu chưa tồn tại
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Đường dẫn tới Tesseract OCR
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        
        # Tạo tài liệu Word
        doc = Document()

        # Chuyển đổi từng trang của PDF thành hình ảnh
        images = convert_from_path(input_path)

        # Duyệt qua từng hình ảnh và trích xuất văn bản
        for i, image in enumerate(images):
            # Sử dụng Tesseract OCR để trích xuất văn bản từ ảnh
            text = pytesseract.image_to_string(image, lang='vie')  # Ngôn ngữ OCR là tiếng Việt
            
            # Ghi văn bản vào tài liệu Word
            doc.add_paragraph(f'Văn bản từ trang {i + 1}:')
            doc.add_paragraph(text)
            doc.add_page_break()
        
        # Lưu tài liệu Word sau khi xử lý xong
        doc.save(output_path)

        # Tạo URL tải file DOCX
        encoded_file_name = quote(output_filename)
        download_url = f"http://127.0.0.1:8000/nguoi_dung/download_docx_converted/{encoded_file_name}"
        
        # Cập nhật thông tin file trong cơ sở dữ liệu
        new_file.converted_doc.name = f'nguoidung/doc/{output_filename}'
        new_file.preview_file = download_url
        new_file.save()
        
        # Lấy dữ liệu JSON của File vừa tạo bằng serializer
        serializer = FileSerializer(new_file)
        
        # Trả về JSON của File cùng với HTTP status code 201 (Created)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    # Nếu file không được gửi trong yêu cầu
    return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)