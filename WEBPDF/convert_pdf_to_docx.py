from pdf2image import convert_from_path
import pytesseract
from docx import Document
from PIL import Image

# Đường dẫn đến tệp PDF cần chuyển đổi
pdf_path = r'file_pdf/1204 KH-HVM_0001.pdf'


# Đường dẫn đến Tesseract OCR (thay đổi theo đường dẫn bạn đã cài đặt)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Tạo tài liệu Word
doc = Document()

# Chuyển đổi từng trang của PDF thành hình ảnh
images = convert_from_path(pdf_path)

# Duyệt qua từng hình ảnh và trích xuất văn bản
for i, image in enumerate(images):
    # Sử dụng Tesseract OCR để trích xuất văn bản từ ảnh
    text = pytesseract.image_to_string(image, lang='vie')  # Thay 'vie' bằng ngôn ngữ khác nếu cần

    # Ghi văn bản vào tệp Word
    doc.add_paragraph(f'Văn bản từ trang {i + 1}:')
    doc.add_paragraph(text)

    # Ngăn cách giữa các trang
    doc.add_page_break()

# Lưu tệp Word sau khi đã xử lý xong
doc.save('output5.docx')

print('Chuyển đổi hoàn thành. Kết quả được lưu vào output4.docx')
